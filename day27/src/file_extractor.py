import os
from pathlib import Path
from typing import Optional, List

import pdfplumber
import docx as docx_lib
import openpyxl


class FileExtractionError(Exception):
    pass


def _extract_pdf(path: Path) -> str:
    text_parts = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

    except Exception as e:
        raise FileExtractionError(f"PDF extraction failed: {e}")
    
    return "\n\n".join(text_parts)


def _extract_docx(path: Path) -> str:
    try:
        doc = docx_lib.Document(str(path))
        

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        table_texts = []
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text])
                if row_text:
                    table_texts.append(row_text)
        
        return "\n\n".join(paragraphs + table_texts)
    except Exception as e:
        raise FileExtractionError(f"DOCX extraction failed: {e}")



def _extract_xlsx(path: Path) -> str:
    try:
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            # Extract values only to avoid formula strings
            for row in ws.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    text_parts.append(row_str)
        return "\n".join(text_parts)
    except Exception as e:
        raise FileExtractionError(f"XLSX extraction failed: {e}")


def _extract_txt(path: Path) -> str:
    encodings = ['utf-8', 'cp1251', 'iso-8859-5']
    
    for encoding in encodings:
        try:
            with open(path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
            
    raise FileExtractionError(f"Could not decode file {path} with supported encodings.")


def extract_text_from_file(file_path: Path) -> str:
    """Extracts text from supported document formats."""
    
    if not file_path.exists():
        raise FileNotFoundError(f"File does not exist: {file_path}")


    ext = file_path.suffix.lower()


    try:
        if ext == '.pdf':
            return _extract_pdf(file_path)
        elif ext in ['.doc', '.docx']:

            return _extract_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            return _extract_xlsx(file_path)

        elif ext == '.txt':
            return _extract_txt(file_path)
        else:
            raise FileExtractionError(f"Unsupported file extension: {ext}")
            
    except FileExtractionError as e:
        raise e
    except Exception as e:
        raise FileExtractionError(f"Unexpected error during extraction: {e}")

